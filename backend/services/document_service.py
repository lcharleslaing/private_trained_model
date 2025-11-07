from pathlib import Path
import hashlib
import json
from typing import List, Dict, Optional
import datetime
import time
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
import numpy as np
import pickle
import os
from dotenv import load_dotenv

# Excel support
try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# OCR imports (optional - only if available)
try:
    import easyocr
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# PDF image extraction (optional)
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

load_dotenv()


class DocumentService:
    def __init__(self, documents_dir: Path = Path("documents"), embeddings_dir: Path = Path("embeddings"), vision_service=None):
        self.documents_dir = documents_dir
        self.embeddings_dir = embeddings_dir
        self.embeddings_dir.mkdir(exist_ok=True)
        
        # Get embedding model from environment
        embedding_model_name = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
        
        # Initialize embedding model (runs locally, no internet needed after first download)
        self.embedding_model = SentenceTransformer(embedding_model_name)
        
        # Initialize OCR reader if available (lazy loading - only when needed)
        self._ocr_reader = None
        self._ocr_enabled = os.getenv("ENABLE_OCR", "true").lower() == "true" and OCR_AVAILABLE
        
        # Vision model support (for understanding drawings)
        self._vision_enabled = os.getenv("ENABLE_VISION", "true").lower() == "true"
        self._vision_service = vision_service  # OllamaService instance for vision
        
        # Load existing embeddings
        self.embeddings_index: Dict[str, List[Dict]] = {}
        self._load_embeddings()
        
        # PERFORMANCE: Cache for loaded chunks (avoids reloading from disk every query)
        self._chunks_cache: Optional[List[Dict]] = None
        self._cache_timestamp: Optional[float] = 0
    
    def _get_ocr_reader(self):
        """Lazy load OCR reader (only initialize when needed)"""
        if not self._ocr_enabled:
            return None
        if self._ocr_reader is None:
            print("Initializing OCR reader (first time may take a moment to download models)...")
            # Initialize with English by default, can add more languages
            self._ocr_reader = easyocr.Reader(['en'], gpu=False)  # gpu=False for CPU-only
        return self._ocr_reader
    
    def _extract_text_with_ocr(self, image_path: Path) -> str:
        """Extract text from an image using OCR"""
        if not self._ocr_enabled:
            return ""
        
        reader = self._get_ocr_reader()
        if reader is None:
            return ""
        
        try:
            results = reader.readtext(str(image_path))
            # Combine all detected text
            text = "\n".join([result[1] for result in results])  # result[1] is the text
            return text
        except Exception as e:
            print(f"OCR error on {image_path.name}: {str(e)}")
            return ""
    
    def _load_embeddings(self):
        """Load existing embeddings from disk"""
        index_file = self.embeddings_dir / "index.json"
        if index_file.exists():
            with open(index_file, "r") as f:
                self.embeddings_index = json.load(f)
    
    def _save_embeddings(self):
        """Save embeddings index to disk"""
        index_file = self.embeddings_dir / "index.json"
        with open(index_file, "w") as f:
            json.dump(self.embeddings_index, f, indent=2)
    
    def _get_document_id(self, file_path: Path) -> str:
        """Generate a unique document ID based on file content"""
        # Use file content hash for better duplicate detection
        try:
            with open(file_path, "rb") as f:
                file_content = f.read()
            return hashlib.md5(file_content).hexdigest()
        except:
            # Fallback to filename + mtime if can't read
            content = f"{file_path.name}_{file_path.stat().st_mtime}"
            return hashlib.md5(content.encode()).hexdigest()
    
    def document_exists(self, file_path: Path) -> bool:
        """Check if a document with the same content already exists"""
        doc_id = self._get_document_id(file_path)
        return doc_id in self.embeddings_index
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file using both text extraction and OCR for images/drawings"""
        text = ""
        
        # Step 1: Try standard text extraction
        try:
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        # Step 2: If very little text extracted, try OCR and Vision on images
        if len(text.strip()) < 50 and PDF2IMAGE_AVAILABLE:
            print(f"PDF '{file_path.name}' has little text. Processing images/drawings...")
            try:
                # Convert PDF pages to images
                images = convert_from_path(str(file_path), dpi=200)  # 200 DPI for good quality
                
                ocr_text = ""
                vision_descriptions = ""
                image_paths = []
                
                # Save images temporarily for processing
                for i, image in enumerate(images):
                    temp_image_path = self.embeddings_dir / f"temp_page_{i}.png"
                    image.save(temp_image_path, "PNG")
                    image_paths.append(temp_image_path)
                
                # Step 2a: Run OCR if enabled
                if self._ocr_enabled:
                    print("Running OCR on images...")
                    for i, temp_image_path in enumerate(image_paths):
                        page_ocr_text = self._extract_text_with_ocr(temp_image_path)
                        if page_ocr_text:
                            ocr_text += f"\n[Page {i+1} - OCR Text]:\n{page_ocr_text}\n"
                
                # Step 2b: Run Vision model if enabled and available
                if self._vision_enabled and self._vision_service:
                    try:
                        if self._vision_service.check_vision_model_available():
                            print("Running vision model to understand drawings...")
                            descriptions = await self._vision_service.describe_images_batch(image_paths)
                            
                            for i, description in enumerate(descriptions):
                                if description:
                                    vision_descriptions += f"\n[Page {i+1} - Vision Model Description]:\n{description}\n"
                            
                            if vision_descriptions:
                                print(f"Vision model generated descriptions for {len([d for d in descriptions if d])} pages.")
                        else:
                            print(f"Vision model '{self._vision_service.vision_model}' not available. Install with: ollama pull {self._vision_service.vision_model}")
                    except Exception as e:
                        print(f"Vision model processing failed: {str(e)}. Continuing with OCR only.")
                
                # Combine OCR and Vision results
                if ocr_text:
                    text += "\n\n[OCR Extracted Text from Images/Drawings]:\n" + ocr_text
                    print(f"OCR extracted {len(ocr_text)} characters from images.")
                
                if vision_descriptions:
                    text += "\n\n[Vision Model Descriptions of Images/Drawings]:\n" + vision_descriptions
                    print(f"Vision model generated {len(vision_descriptions)} characters of descriptions.")
                
                # Clean up temp files
                for temp_image_path in image_paths:
                    if temp_image_path.exists():
                        temp_image_path.unlink()
                
                if not ocr_text and not vision_descriptions:
                    print("Neither OCR nor vision model extracted content from images.")
                    
            except Exception as e:
                print(f"Image processing failed: {str(e)}. Continuing with extracted text only.")
        
        # Warn if still very little text
        if len(text.strip()) < 50:
            print(f"Warning: PDF '{file_path.name}' has minimal text ({len(text.strip())} chars). "
                  f"May contain only images/drawings without extractable text.")
        
        return text
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX or DOC file
        
        Note: .doc files (older Word format) may not be fully supported.
        python-docx primarily works with .docx files. For .doc files, 
        consider converting to .docx first for best results.
        """
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n[Table]:\n" + "\n".join(table_text))
            
            result = "\n".join(text_parts)
            
            # Warn if .doc file and little content extracted
            if file_path.suffix.lower() == ".doc" and len(result.strip()) < 50:
                print(f"Warning: .doc file '{file_path.name}' may not be fully supported. "
                      f"Consider converting to .docx for better results.")
            
            return result
        except Exception as e:
            if file_path.suffix.lower() == ".doc":
                raise Exception(f"Error reading .doc file: {str(e)}. "
                              f".doc files (older Word format) may not be fully supported. "
                              f"Consider converting to .docx format.")
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_text_from_xlsx(self, file_path: Path) -> str:
        """Extract text from Excel file (.xlsx, .xls)
        
        Extracts data from all sheets, preserving structure.
        Limits to 1000 rows per sheet to avoid processing huge files.
        """
        if not EXCEL_AVAILABLE:
            raise Exception("Excel support not available. Install openpyxl and pandas: pip install openpyxl pandas")
        
        try:
            text_parts = []
            
            # Read Excel file
            excel_file = pd.ExcelFile(file_path)
            
            print(f"Processing Excel file with {len(excel_file.sheet_names)} sheet(s)...")
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                text_parts.append(f"\n[Sheet: {sheet_name}]\n")
                
                # Read sheet
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Convert to text representation
                if not df.empty:
                    # Include headers
                    headers = " | ".join([str(col) for col in df.columns])
                    text_parts.append(f"Headers: {headers}\n")
                    
                    # Include data rows (limit to avoid huge files)
                    max_rows = 1000  # Limit rows per sheet
                    rows_to_process = min(len(df), max_rows)
                    
                    for idx, row in df.head(max_rows).iterrows():
                        row_data = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                        text_parts.append(f"Row {idx + 1}: {row_data}")
                    
                    if len(df) > max_rows:
                        text_parts.append(f"\n... ({len(df) - max_rows} more rows not shown)")
                    
                    print(f"  Sheet '{sheet_name}': {rows_to_process} rows processed")
                else:
                    text_parts.append("(Empty sheet)")
                    print(f"  Sheet '{sheet_name}': Empty")
            
            result = "\n".join(text_parts)
            print(f"Excel file processed: {len(result)} characters extracted")
            return result
        except Exception as e:
            raise Exception(f"Error reading Excel file: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    async def _extract_text_from_image(self, file_path: Path) -> str:
        """Extract text and description from image file (JPG, PNG, etc.)
        
        Uses both OCR (for text) and Vision Model (for understanding)
        """
        text_parts = []
        
        # Step 1: Run OCR if enabled
        if self._ocr_enabled:
            print(f"Running OCR on image: {file_path.name}")
            ocr_text = self._extract_text_with_ocr(file_path)
            if ocr_text:
                text_parts.append(f"[OCR Extracted Text]:\n{ocr_text}")
                print(f"OCR extracted {len(ocr_text)} characters from image.")
        
        # Step 2: Run Vision Model if enabled and available
        if self._vision_enabled and self._vision_service:
            try:
                if self._vision_service.check_vision_model_available():
                    print(f"Running vision model on image: {file_path.name}")
                    description = await self._vision_service.describe_image(file_path)
                    if description:
                        text_parts.append(f"\n[Vision Model Description]:\n{description}")
                        print(f"Vision model generated {len(description)} characters of description.")
                else:
                    print(f"Vision model '{self._vision_service.vision_model}' not available. Install with: ollama pull {self._vision_service.vision_model}")
            except Exception as e:
                print(f"Vision model processing failed: {str(e)}. Continuing with OCR only.")
        
        if not text_parts:
            return f"[Image file: {file_path.name}]\nNo text or description could be extracted from this image."
        
        return "\n".join(text_parts)
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats (async for image processing)"""
        suffix = file_path.suffix.lower()
        
        # Image formats
        image_formats = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"]
        
        if suffix == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._extract_text_from_docx(file_path)
        elif suffix in [".xlsx", ".xls"]:
            return self._extract_text_from_xlsx(file_path)
        elif suffix == ".txt":
            return self._extract_text_from_txt(file_path)
        elif suffix in image_formats:
            return await self._extract_text_from_image(file_path)
        else:
            raise Exception(f"Unsupported file format: {suffix}. Supported formats: PDF, DOCX, DOC, XLSX, XLS, TXT, JPG, PNG, GIF, BMP, TIFF, WEBP")
    
    def _chunk_text(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[str]:
        """Split text into overlapping chunks"""
        # Get chunk settings from environment
        if chunk_size is None:
            chunk_size = int(os.getenv("CHUNK_SIZE", "500"))
        if overlap is None:
            overlap = int(os.getenv("CHUNK_OVERLAP", "50"))
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
        
        return chunks
    
    async def process_document(self, file_path: Path, force_reprocess: bool = False) -> Dict:
        """Process a document: extract text, chunk it, and create embeddings
        
        Args:
            file_path: Path to the document file
            force_reprocess: If True, reprocess even if document already exists
        """
        # Generate document ID
        document_id = self._get_document_id(file_path)
        
        # Check if document already exists
        if document_id in self.embeddings_index and not force_reprocess:
            existing_doc = self.embeddings_index[document_id]
            return {
                "document_id": document_id,
                "chunks": existing_doc["chunks"],
                "chunks_data": [],
                "message": "Document already exists in database",
                "already_exists": True
            }
        
        # Extract text (async for images)
        text = await self._extract_text(file_path)
        
        if not text.strip():
            raise Exception("Document appears to be empty")
        
        # Chunk the text
        chunks = self._chunk_text(text)
        
        # Create embeddings for each chunk
        embeddings = self.embedding_model.encode(chunks, show_progress_bar=False)
        
        # Store chunks and embeddings
        document_chunks = []
        chunks_metadata = {}
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{document_id}_chunk_{i}"
            chunk_data = {
                "chunk_id": chunk_id,
                "content": chunk,
                "document_id": document_id,
                "source": file_path.name,
                "chunk_index": i
            }
            document_chunks.append(chunk_data)
            
            # Store chunk metadata
            chunks_metadata[chunk_id] = {
                "content": chunk,
                "document_id": document_id,
                "source": file_path.name,
                "chunk_index": i
            }
            
            # Store embedding
            embedding_file = self.embeddings_dir / f"{chunk_id}.pkl"
            with open(embedding_file, "wb") as f:
                pickle.dump(embedding, f)
        
        # Store chunks metadata
        chunks_file = self.embeddings_dir / f"{document_id}_chunks.json"
        with open(chunks_file, "w", encoding="utf-8") as f:
            json.dump(chunks_metadata, f, indent=2, ensure_ascii=False)
        
        # Update index with metadata
        self.embeddings_index[document_id] = {
            "filename": file_path.name,
            "chunks": len(document_chunks),
            "chunk_ids": [chunk["chunk_id"] for chunk in document_chunks],
            "uploaded_at": datetime.datetime.now().isoformat(),
            "file_size": file_path.stat().st_size,
            "file_path": str(file_path)
        }
        
        # Save index
        self._save_embeddings()
        
        # Invalidate cache since we added new chunks
        self.invalidate_cache()
        
        return {
            "document_id": document_id,
            "chunks": len(document_chunks),
            "chunks_data": document_chunks,
            "already_exists": False,
            "message": "Document processed and added to database"
        }
    
    def list_documents(self) -> List[Dict]:
        """List all processed documents with metadata"""
        documents = []
        for doc_id, doc_info in self.embeddings_index.items():
            doc_data = {
                "document_id": doc_id,
                "filename": doc_info["filename"],
                "chunks": doc_info["chunks"],
                "uploaded_at": doc_info.get("uploaded_at", "Unknown"),
                "file_size": doc_info.get("file_size", 0)
            }
            documents.append(doc_data)
        # Sort by upload date (newest first)
        documents.sort(key=lambda x: x.get("uploaded_at", ""), reverse=True)
        return documents
    
    def get_document_stats(self) -> Dict:
        """Get statistics about the document database"""
        total_docs = len(self.embeddings_index)
        total_chunks = sum(doc_info.get("chunks", 0) for doc_info in self.embeddings_index.values())
        total_size = sum(doc_info.get("file_size", 0) for doc_info in self.embeddings_index.values())
        
        return {
            "total_documents": total_docs,
            "total_chunks": total_chunks,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2)
        }
    
    def delete_document(self, document_id: str):
        """Delete a document and its embeddings"""
        if document_id not in self.embeddings_index:
            raise Exception("Document not found")
        
        # Delete embedding files
        chunk_ids = self.embeddings_index[document_id]["chunk_ids"]
        for chunk_id in chunk_ids:
            embedding_file = self.embeddings_dir / f"{chunk_id}.pkl"
            if embedding_file.exists():
                embedding_file.unlink()
        
        # Delete chunks metadata file
        chunks_file = self.embeddings_dir / f"{document_id}_chunks.json"
        if chunks_file.exists():
            chunks_file.unlink()
        
        # Remove from index
        del self.embeddings_index[document_id]
        self._save_embeddings()
        
        # Invalidate cache since we deleted chunks
        self.invalidate_cache()
    
    def get_all_chunks(self) -> List[Dict]:
        """Get all document chunks with their embeddings (loads from disk)"""
        all_chunks = []
        for doc_id, doc_info in self.embeddings_index.items():
            # Load chunks metadata for this document
            chunks_file = self.embeddings_dir / f"{doc_id}_chunks.json"
            chunks_metadata = {}
            if chunks_file.exists():
                with open(chunks_file, "r", encoding="utf-8") as f:
                    chunks_metadata = json.load(f)
            
            for chunk_id in doc_info["chunk_ids"]:
                embedding_file = self.embeddings_dir / f"{chunk_id}.pkl"
                if embedding_file.exists():
                    with open(embedding_file, "rb") as f:
                        embedding = pickle.load(f)
                    
                    # Get chunk metadata
                    chunk_meta = chunks_metadata.get(chunk_id, {})
                    
                    all_chunks.append({
                        "chunk_id": chunk_id,
                        "document_id": doc_id,
                        "source": doc_info["filename"],
                        "content": chunk_meta.get("content", ""),
                        "embedding": np.array(embedding)  # Convert to numpy array for faster operations
                    })
        return all_chunks
    
    def get_all_chunks_cached(self) -> List[Dict]:
        """Get all chunks with caching for performance"""
        current_time = time.time()
        cache_ttl = 300  # Cache for 5 minutes
        
        # Check if cache is valid
        if (self._chunks_cache is not None and 
            (current_time - self._cache_timestamp) < cache_ttl):
            return self._chunks_cache
        
        # Load and cache
        self._chunks_cache = self.get_all_chunks()
        self._cache_timestamp = current_time
        return self._chunks_cache
    
    def invalidate_cache(self):
        """Invalidate the chunks cache (call after adding/deleting documents)"""
        self._chunks_cache = None
        self._cache_timestamp = 0
    
    async def reindex_all(self) -> Dict:
        """Reindex all documents in the documents directory"""
        count = 0
        supported_formats = [".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", 
                            ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"]
        for file_path in self.documents_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_formats:
                try:
                    await self.process_document(file_path, force_reprocess=True)
                    count += 1
                except Exception as e:
                    print(f"Error processing {file_path.name}: {str(e)}")
        return {"count": count}

